import scrapy
import json
from scrapy_splash import SplashRequest
import urllib.request
from docx import Document
import os.path


class KnowledgesSpider(scrapy.Spider):
    name = "knowledges"
    allowed_domains = ["www.sanfranroaster.com"]

    script = """
      function main(splash, args)
        assert(splash:go(args.url))
        assert(splash:wait(1))
        return splash:html()
      end
    """

    def start_requests(self):
        f = open(os.path.dirname(__file__) + "/../links.json")
        links = json.load(f)

        for link in links:
            print(link["link"])
            yield SplashRequest(
                url=link["link"],
                endpoint="execute",
                callback=self.parse,
                args={"lua_source": self.script},
                meta={
                    "mod": link["model"],
                    "des": link["description"],
                    "lin": link["link"],
                },
            )

    def parse(self, response):
        nopage = response.xpath("//div[@class='kb__404-section']").get()
        model = response.request.meta["mod"]
        description = response.request.meta["des"]
        link = response.request.meta["lin"]

        document = Document()

        if nopage is None:
            title = response.xpath("//span[@id='hs_cos_wrapper_name']/text()").get()
            subtitle = response.xpath("//h2/text()").get()
            items = response.xpath(""" //div[@class='kb-article tinymce-content']/p """)
            yield {"title": title, "subtitle": subtitle, "items": items}

            document.add_heading(title, 0)
            document.add_heading(subtitle, level=2)
            document.add_paragraph(
                f"メタ情報 ::: モデル：{model}, 詳細：{description}, リンク：{link}"
            )

            for item in items:
                image = item.xpath(".//img/@src").get()
                if image is None:
                    text = item.xpath(".//text()").get()
                    document.add_paragraph(text)
                else:
                    filename = image.split("/")[-1]
                    urllib.request.urlretrieve(image, filename)
                    document.add_picture(filename)

        else:
            document.add_heading(description, 0)
            document.add_paragraph(
                f"メタ情報 ::: モデル：{model}, 詳細：{description}, リンク：{link}"
            )
            document.add_paragraph("リンクが壊れてます。")

        document.save(f"Sanfranciscan knowledgebase {description}.docx")
